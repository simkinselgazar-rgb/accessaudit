"""Export ACR report data to a DOCX document using python-docx."""
from __future__ import annotations

import logging
import os
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from models import TestResult, ConformanceLevel, ReviewMeta

logger = logging.getLogger(__name__)


def _resolve_evidence_path(screenshot_path: str, output_path: Path) -> Path | None:
    """Resolve a finding's screenshot_path to an absolute file path on disk.

    The screenshot_path field in findings can be either:
      - absolute (already a full path)
      - relative to the review directory ('captures/full_page.png')
      - relative to the captures directory ('full_page.png')

    The DOCX is written to ``<review_dir>/report/acr_report.docx``, so the
    review directory is the grandparent of output_path. Returns None if no
    valid file exists at any candidate location.
    """
    if not screenshot_path:
        return None
    candidate = Path(screenshot_path)
    if candidate.is_absolute() and candidate.exists():
        return candidate
    review_dir = output_path.parent.parent
    for base in (review_dir, review_dir / "captures"):
        try_path = (base / screenshot_path).resolve()
        if try_path.exists() and try_path.is_file():
            return try_path
    return None


def _embed_screenshot(doc: Document, screenshot_path: str | None, output_path: Path, *, width_in: float = 5.5) -> bool:
    """Embed a screenshot into the DOCX as a centered image. Returns True if added."""
    if not screenshot_path:
        return False
    resolved = _resolve_evidence_path(screenshot_path, output_path)
    if resolved is None:
        return False
    try:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(str(resolved), width=Inches(width_in))
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = caption.add_run(f"Evidence: {resolved.name}")
        cap_run.italic = True
        cap_run.font.size = Pt(9)
        cap_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        return True
    except Exception as exc:
        logger.warning("Could not embed screenshot %s: %s", resolved, exc)
        return False

# ---------------------------------------------------------------------------
# Colour palette (matching the HTML templates)
# ---------------------------------------------------------------------------
_COLORS: dict[str, RGBColor] = {
    "Supports": RGBColor(0x2E, 0x7D, 0x32),
    "Partially Supports": RGBColor(0xF5, 0x7F, 0x17),
    "Does Not Support": RGBColor(0xC6, 0x28, 0x28),
    "Not Applicable": RGBColor(0x75, 0x75, 0x75),
    "Not Evaluated": RGBColor(0x7B, 0x1F, 0xA2),
}

_BG_COLORS: dict[str, str] = {
    "Supports": "E8F5E9",
    "Partially Supports": "FFF8E1",
    "Does Not Support": "FFEBEE",
    "Not Applicable": "F5F5F5",
    "Not Evaluated": "F3E5F5",
}

_SEVERITY_COLORS: dict[str, RGBColor] = {
    "high": RGBColor(0xC6, 0x28, 0x28),
    "medium": RGBColor(0xF5, 0x7F, 0x17),
    "low": RGBColor(0x2E, 0x7D, 0x32),
    "info": RGBColor(0x75, 0x75, 0x75),
}


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

from report.utils import (
    FPC_MAPPING,
    FPC_NAMES,
    ResultProxy as _ResultProxy,
    build_fpc_rows as _build_fpc_rows,
    conformance_str as _conformance_str,
    criterion_sort_key as _sort_key,
    wrap_results as _wrap_results,
)


def _set_cell_bg(cell, hex_color: str) -> None:
    """Apply a background colour to a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shading_el = shading.find(qn("w:shd"))
    if shading_el is None:
        from lxml import etree
        shading_el = etree.SubElement(shading, qn("w:shd"))
    shading_el.set(qn("w:fill"), hex_color)
    shading_el.set(qn("w:val"), "clear")


def _add_conformance_cell(row, col_idx: int, value: str) -> None:
    """Write the conformance value with appropriate colour."""
    cell = row.cells[col_idx]
    cell.text = ""
    run = cell.paragraphs[0].add_run(value)
    run.bold = True
    color = _COLORS.get(value)
    if color:
        run.font.color.rgb = color
    bg = _BG_COLORS.get(value)
    if bg:
        _set_cell_bg(cell, bg)


def _style_header_row(table) -> None:
    """Make the first row bold with a dark background."""
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_bg(cell, "37474F")


def _add_criteria_table(
    doc: Document,
    heading: str,
    rows_data: list[dict[str, Any]],
) -> None:
    """Add a criteria conformance table for a WCAG level (internal mode)."""
    if not rows_data:
        return
    doc.add_heading(heading, level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = "Criteria"
    hdr[1].text = "Name"
    hdr[2].text = "Conformance Level"
    hdr[3].text = "Remarks and Explanations"
    _style_header_row(table)

    rows_data_sorted = sorted(rows_data, key=lambda r: _sort_key(r["criterion_id"]))
    for item in rows_data_sorted:
        row = table.add_row()
        row.cells[0].text = item["criterion_id"]
        row.cells[1].text = item.get("criterion_name", "")
        conformance = item.get("conformance_display", item.get("conformance_level", ""))
        _add_conformance_cell(row, 2, conformance)
        row.cells[3].text = item.get("summary", "")

    doc.add_paragraph()  # spacer


# ---------------------------------------------------------------------------
# Section 508 cross-references for each WCAG SC (VPAT 2.5 Rev template)
# ---------------------------------------------------------------------------
_SEC508_REFS: dict[str, str] = {
    "1.1.1": "501 (Web)(Software)\n504.2 (Authoring Tool)\n602.3 (Support Docs)",
    "1.2.1": "501 (Web)(Software)\n504.2 (Authoring Tool)\n602.3 (Support Docs)",
    "1.2.2": "501 (Web)(Software)\n504.2 (Authoring Tool)\n602.3 (Support Docs)",
    "1.2.3": "501 (Web)(Software)\n504.2 (Authoring Tool)\n602.3 (Support Docs)",
    "1.2.4": "501 (Web)(Software)\n504.2 (Authoring Tool)\n602.3 (Support Docs)",
    "1.2.5": "501 (Web)(Software)\n504.2 (Authoring Tool)\n602.3 (Support Docs)",
    "1.3.1": "501 (Web)(Software)\n504.2 (Authoring Tool)\n602.3 (Support Docs)",
    "1.3.2": "501 (Web)(Software)\n504.2 (Authoring Tool)\n602.3 (Support Docs)",
    "1.3.3": "501 (Web)(Software)\n504.2 (Authoring Tool)\n602.3 (Support Docs)",
    "1.3.4": "501 (Web)(Software)\n504.2 (Authoring Tool)\n602.3 (Support Docs)",
    "1.3.5": "501 (Web)(Software)\n504.2 (Authoring Tool)\n602.3 (Support Docs)",
    "1.4.1": "501 (Web)(Software)\n504.2 (Authoring Tool)\n602.3 (Support Docs)",
    "1.4.2": "501 (Web)(Software)\n504.2 (Authoring Tool)\n602.3 (Support Docs)",
    "1.4.3": "501 (Web)(Software)\n504.2 (Authoring Tool)\n602.3 (Support Docs)",
    "1.4.4": "501 (Web)(Software)\n504.2 (Authoring Tool)\n602.3 (Support Docs)",
    "1.4.5": "501 (Web)(Software)\n504.2 (Authoring Tool)\n602.3 (Support Docs)",
    "1.4.10": "501 (Web)(Software)\n504.2 (Authoring Tool)\n602.3 (Support Docs)",
    "1.4.11": "501 (Web)(Software)\n504.2 (Authoring Tool)\n602.3 (Support Docs)",
    "1.4.12": "501 (Web)(Software)\n504.2 (Authoring Tool)\n602.3 (Support Docs)",
    "1.4.13": "501 (Web)(Software)\n504.2 (Authoring Tool)\n602.3 (Support Docs)",
    "2.1.1": "501 (Web)(Software)\n504.2 (Authoring Tool)\n602.3 (Support Docs)",
    "2.1.2": "501 (Web)(Software)\n504.2 (Authoring Tool)\n602.3 (Support Docs)",
    "2.1.4": "501 (Web)(Software)\n504.2 (Authoring Tool)\n602.3 (Support Docs)",
    "2.2.1": "501 (Web)(Software)\n504.2 (Authoring Tool)\n602.3 (Support Docs)",
    "2.2.2": "501 (Web)(Software)\n504.2 (Authoring Tool)\n602.3 (Support Docs)",
    "2.3.1": "501 (Web)(Software)\n504.2 (Authoring Tool)\n602.3 (Support Docs)",
    "2.4.1": "501 (Web)(Software) \u2013 Does not apply to non-web software\n504.2 (Authoring Tool)\n602.3 (Support Docs)",
    "2.4.2": "501 (Web)(Software)\n504.2 (Authoring Tool)\n602.3 (Support Docs)",
    "2.4.3": "501 (Web)(Software)\n504.2 (Authoring Tool)\n602.3 (Support Docs)",
    "2.4.4": "501 (Web)(Software)\n504.2 (Authoring Tool)\n602.3 (Support Docs)",
    "2.4.5": "501 (Web)(Software) \u2013 Does not apply to non-web software\n504.2 (Authoring Tool)\n602.3 (Support Docs)",
    "2.4.6": "501 (Web)(Software)\n504.2 (Authoring Tool)\n602.3 (Support Docs)",
    "2.4.7": "501 (Web)(Software)\n504.2 (Authoring Tool)\n602.3 (Support Docs)",
    "2.5.1": "501 (Web)(Software)\n504.2 (Authoring Tool)\n602.3 (Support Docs)",
    "2.5.2": "501 (Web)(Software)\n504.2 (Authoring Tool)\n602.3 (Support Docs)",
    "2.5.3": "501 (Web)(Software)\n504.2 (Authoring Tool)\n602.3 (Support Docs)",
    "2.5.4": "501 (Web)(Software)\n504.2 (Authoring Tool)\n602.3 (Support Docs)",
    "3.1.1": "501 (Web)(Software)\n504.2 (Authoring Tool)\n602.3 (Support Docs)",
    "3.1.2": "501 (Web)(Software)\n504.2 (Authoring Tool)\n602.3 (Support Docs)",
    "3.2.1": "501 (Web)(Software)\n504.2 (Authoring Tool)\n602.3 (Support Docs)",
    "3.2.2": "501 (Web)(Software)\n504.2 (Authoring Tool)\n602.3 (Support Docs)",
    "3.2.3": "501 (Web)(Software) \u2013 Does not apply to non-web software\n504.2 (Authoring Tool)\n602.3 (Support Docs)",
    "3.2.4": "501 (Web)(Software) \u2013 Does not apply to non-web software\n504.2 (Authoring Tool)\n602.3 (Support Docs)",
    "3.3.1": "501 (Web)(Software)\n504.2 (Authoring Tool)\n602.3 (Support Docs)",
    "3.3.2": "501 (Web)(Software)\n504.2 (Authoring Tool)\n602.3 (Support Docs)",
    "3.3.3": "501 (Web)(Software)\n504.2 (Authoring Tool)\n602.3 (Support Docs)",
    "3.3.4": "501 (Web)(Software)\n504.2 (Authoring Tool)\n602.3 (Support Docs)",
    "4.1.1": "501 (Web)(Software)\n504.2 (Authoring Tool)\n602.3 (Support Docs)",
    "4.1.2": "501 (Web)(Software)\n504.2 (Authoring Tool)\n602.3 (Support Docs)",
    "4.1.3": "501 (Web)(Software)\n504.2 (Authoring Tool)\n602.3 (Support Docs)",
}

# VPAT terms definitions
_VPAT_TERMS: list[tuple[str, str]] = [
    ("Supports", "The functionality of the product has at least one method that meets the criterion without known defects or meets with equivalent facilitation."),
    ("Partially Supports", "Some functionality of the product does not meet the criterion."),
    ("Does Not Support", "The majority of product functionality does not meet the criterion."),
    ("Not Applicable", "The criterion is not relevant to the product."),
    ("Not Evaluated", "The product has not been evaluated against the criterion. This can only be used in WCAG Level AAA criteria."),
]




# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def _sanitize_eval_methods(text: str) -> str:
    """Strip AI language for client-facing reports."""
    return (
        text
        .replace("AI-powered ", "")
        .replace("AI-assisted ", "")
        .replace(", AI-powered visual analysis,", ", visual analysis,")
    )


def _add_vpat_criteria_cell(cell, criterion_id: str, criterion_name: str,
                            level: str) -> None:
    """Build the VPAT-format criteria cell with 508 cross-references."""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(f"{criterion_id} {criterion_name} (Level {level})")
    run.bold = True
    run.font.size = Pt(9)

    ref = _SEC508_REFS.get(criterion_id)
    if ref:
        lines = f"\nAlso applies to:\nRevised Section 508\n{ref}"
        ref_run = p.add_run(lines)
        ref_run.font.size = Pt(8)
        ref_run.font.color.rgb = RGBColor(0x54, 0x54, 0x54)


def _add_vpat_wcag_table(doc: Document, heading: str, rows_data: list[dict],
                         level: int = 3) -> None:
    """Add a 3-column VPAT-format WCAG criteria table."""
    if not rows_data:
        return
    doc.add_heading(heading, level=level)
    doc.add_paragraph("Notes:")

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = "Criteria"
    hdr[1].text = "Conformance Level"
    hdr[2].text = "Remarks and Explanations"
    _style_header_row(table)

    rows_sorted = sorted(rows_data, key=lambda r: _sort_key(r["criterion_id"]))
    for item in rows_sorted:
        row = table.add_row()
        _add_vpat_criteria_cell(
            row.cells[0],
            item["criterion_id"],
            item.get("criterion_name", ""),
            item.get("level", "A"),
        )
        conformance = item.get("conformance_display", item.get("conformance_level", ""))
        cell_conf = row.cells[1]
        cell_conf.text = ""
        conf_run = cell_conf.paragraphs[0].add_run(f"Web: {conformance}")
        conf_run.bold = True
        color = _COLORS.get(conformance)
        if color:
            conf_run.font.color.rgb = color
        bg = _BG_COLORS.get(conformance)
        if bg:
            _set_cell_bg(cell_conf, bg)

        summary = item.get("summary", "")
        row.cells[2].text = f"Web: {summary}" if summary else "Web: "

    doc.add_paragraph()


def _export_vpat_client(results, meta, output_path: Path) -> str:
    """Generate a VPAT 2.5 Rev Section 508 format DOCX (client-facing)."""
    meta_dict = meta.to_dict()
    company = meta_dict.get("company_name", "") or "[Company]"
    product = meta_dict.get("product_name", "")
    wcag_ver = meta_dict.get("wcag_version", "2.2")
    coverage = meta_dict.get("coverage_level", "AA")

    doc = Document()

    # ── Report Title (Heading 1) ──────────────────────────────────────────
    doc.add_heading(f"{company} Accessibility Conformance Report", level=1)
    doc.add_paragraph("Revised Section 508 Edition")
    p_vpat = doc.add_paragraph()
    r_vpat = p_vpat.add_run("(Based on VPAT\u00ae Version 2.5Rev)")
    r_vpat.font.size = Pt(10)

    doc.add_paragraph()

    # ── Header fields ─────────────────────────────────────────────────────
    report_date = meta_dict.get("created_at", "")
    product_desc = str(meta_dict.get("product_description", "")).strip()
    contact_name = str(meta_dict.get("contact_name", "")).strip()
    contact_email = str(meta_dict.get("contact_email", "")).strip()
    contact_info = ", ".join(filter(None, [contact_name, contact_email]))
    notes = str(meta_dict.get("notes", "")).strip()
    eval_methods = str(meta_dict.get("evaluation_methods", "")).strip()
    eval_methods = _sanitize_eval_methods(eval_methods)

    header_fields = [
        ("Name of Product/Version", product),
        ("Report Date", report_date),
        ("Product Description", product_desc),
        ("Contact Information", contact_info),
        ("Notes", notes),
        ("Evaluation Methods Used", eval_methods),
    ]
    for label, value in header_fields:
        p = doc.add_paragraph()
        run_label = p.add_run(f"{label}: ")
        run_label.bold = True
        p.add_run(value)

    # ── Applicable Standards/Guidelines heading + table ───────────────────
    p_asg = doc.add_paragraph()
    run_asg = p_asg.add_run("Applicable Standards/Guidelines")
    run_asg.bold = True
    doc.add_paragraph(
        "This report covers the degree of conformance for the following "
        "accessibility standards/guidelines:"
    )

    include_aaa = coverage.upper() == "AAA"
    sg_table = doc.add_table(rows=3, cols=2)
    sg_table.style = "Table Grid"
    sg_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sg_hdr = sg_table.rows[0].cells
    sg_hdr[0].text = "Standard/Guideline"
    sg_hdr[1].text = "Included In Report"
    _style_header_row(sg_table)

    sg_table.rows[1].cells[0].text = (
        f"Web Content Accessibility Guidelines {wcag_ver}"
    )
    aaa_yn = "Yes" if include_aaa else "No"
    sg_table.rows[1].cells[1].text = (
        f"Level A (Yes)\nLevel AA (Yes)\nLevel AAA ({aaa_yn})"
    )
    sg_table.rows[2].cells[0].text = (
        "Revised Section 508 standards published January 18, 2017 "
        "and corrected January 22, 2018"
    )
    sg_table.rows[2].cells[1].text = "(Yes)"
    doc.add_paragraph()

    # ── Terms ─────────────────────────────────────────────────────────────
    doc.add_heading("Terms", level=2)
    doc.add_paragraph(
        "The terms used in the Conformance Level information are defined "
        "as follows:"
    )
    for term, defn in _VPAT_TERMS:
        p = doc.add_paragraph()
        run_t = p.add_run(f"{term}: ")
        run_t.bold = True
        p.add_run(defn)
    doc.add_paragraph()

    # ── WCAG Report ───────────────────────────────────────────────────────
    doc.add_heading(f"WCAG {wcag_ver} Report", level=2)
    doc.add_paragraph(
        "Tables 1 and 2 also document conformance with Revised Section 508:"
    )
    doc.add_paragraph(
        "Chapter 5 \u2013 501.1 Scope, 504.2 Content Creation or Editing"
    )
    doc.add_paragraph(
        "Chapter 6 \u2013 602.3 Electronic Support Documentation"
    )
    doc.add_paragraph(
        "Note: When reporting on conformance with the WCAG 2.x Success "
        "Criteria, they are scoped for full pages, complete processes, and "
        "accessibility-supported ways of using technology as documented in "
        "the WCAG 2.x Conformance Requirements."
    )

    # Bucket results by level
    buckets: dict[str, list[dict]] = {"A": [], "AA": [], "AAA": []}
    for r in results:
        key = r.level.upper().strip()
        if key not in buckets:
            key = "A"
        row = r.to_dict()
        row["conformance_display"] = _conformance_str(r.conformance_level)
        row["level"] = r.level.upper().strip()
        buckets[key].append(row)

    _add_vpat_wcag_table(doc, "Table 1: Success Criteria, Level A", buckets["A"])
    _add_vpat_wcag_table(doc, "Table 2: Success Criteria, Level AA", buckets["AA"])
    if buckets["AAA"]:
        _add_vpat_wcag_table(doc, "Table 3: Success Criteria, Level AAA", buckets["AAA"])

    # ── Revised Section 508 Report ────────────────────────────────────────
    doc.add_heading("Revised Section 508 Report", level=2)
    doc.add_paragraph("Notes:")

    # Chapter 3: FPC
    doc.add_heading("Chapter 3: Functional Performance Criteria (FPC)", level=3)
    doc.add_paragraph("Notes:")
    fpc_rows = _build_fpc_rows(results)
    fpc_table = doc.add_table(rows=1, cols=3)
    fpc_table.style = "Table Grid"
    fpc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    fpc_hdr = fpc_table.rows[0].cells
    fpc_hdr[0].text = "Criteria"
    fpc_hdr[1].text = "Conformance Level"
    fpc_hdr[2].text = "Remarks and Explanations"
    _style_header_row(fpc_table)
    for item in fpc_rows:
        row = fpc_table.add_row()
        row.cells[0].text = f"{item['code']} {item['name']}"
        conformance = item["conformance"]
        cell_conf = row.cells[1]
        cell_conf.text = ""
        conf_run = cell_conf.paragraphs[0].add_run(conformance)
        conf_run.bold = True
        color = _COLORS.get(conformance)
        if color:
            conf_run.font.color.rgb = color
        bg = _BG_COLORS.get(conformance)
        if bg:
            _set_cell_bg(cell_conf, bg)
        row.cells[2].text = item.get("remark", "")
    doc.add_paragraph()

    # ── Legal Disclaimer ──────────────────────────────────────────────────
    doc.add_heading("Legal Disclaimer", level=2)
    disclaimer_para = doc.add_paragraph()
    disclaimer_run = disclaimer_para.add_run(
        "This report is provided for informational purposes only and does "
        "not constitute legal advice. Accessibility conformance was evaluated "
        "based on automated and manual testing methods and may not reflect "
        "all potential accessibility barriers. No warranty, express or "
        "implied, is made regarding the completeness or accuracy of this "
        "report."
    )
    disclaimer_run.italic = True
    disclaimer_run.font.size = Pt(9)
    disclaimer_run.font.color.rgb = RGBColor(0x75, 0x75, 0x75)

    # ── Appendix: Detailed Findings ───────────────────────────────────────
    all_findings = []
    for r in results:
        for f in r.findings:
            entry = f.to_dict() if hasattr(f, "to_dict") else dict(f)
            entry["criterion_id"] = r.criterion_id
            entry["criterion_name"] = r.criterion_name
            all_findings.append(entry)

    if all_findings:
        sev_order = {"high": 0, "medium": 1, "low": 2, "info": 3}
        all_findings.sort(key=lambda x: (
            sev_order.get(x.get("severity", "info"), 4),
            x.get("criterion_id", ""),
        ))
        doc.add_heading("Appendix: Detailed Findings", level=2)
        for finding in all_findings:
            heading_text = (
                f"{finding.get('criterion_id', '')} - "
                f"{finding.get('issue', 'Finding')}"
            )
            doc.add_heading(heading_text, level=3)
            severity = finding.get("severity", "info")
            sev_color = _SEVERITY_COLORS.get(severity, RGBColor(0x75, 0x75, 0x75))
            sev_para = doc.add_paragraph()
            sev_run = sev_para.add_run(f"Severity: {severity.upper()}")
            sev_run.bold = True
            sev_run.font.color.rgb = sev_color

            if finding.get("element"):
                doc.add_paragraph(f"Element: {finding['element']}")
            if finding.get("location"):
                doc.add_paragraph(f"Location: {finding['location']}")
            if finding.get("css_selector"):
                doc.add_paragraph(f"Selector: {finding['css_selector']}")
            if finding.get("impact"):
                doc.add_paragraph(f"Impact: {finding['impact']}")
            if finding.get("recommendation"):
                doc.add_paragraph(f"Recommendation: {finding['recommendation']}")
            # Embed the screenshot evidence so the client can see exactly
            # what the finding refers to. _embed_screenshot resolves the
            # path relative to the review directory and skips silently if
            # the file is missing.
            _embed_screenshot(doc, finding.get("screenshot_path"), output_path)
            doc.add_paragraph()

    doc.save(str(output_path))
    logger.info("VPAT DOCX report (client mode) written to %s", output_path)
    return str(output_path)


def export_docx(
    results: list[TestResult],
    meta: ReviewMeta,
    output_path: str | Path,
    client_mode: bool = False,
) -> str:
    """Generate a DOCX accessibility conformance report.

    Parameters
    ----------
    results : list[TestResult]
        Evaluated criteria results (may also be plain dicts).
    meta : ReviewMeta
        Review metadata.
    output_path : str | Path
        Destination file path for the .docx file.
    client_mode : bool
        When True, produce a VPAT 2.5 Rev Section 508 format report
        suitable for external delivery.

    Returns
    -------
    str
        Absolute path of the written file.
    """
    output_path = Path(output_path).resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)

    # Normalize: accept both TestResult objects and plain dicts
    results = _wrap_results(results)

    if client_mode:
        return _export_vpat_client(results, meta, output_path)

    # ── Internal mode (unchanged) ─────────────────────────────────────────
    report_format = getattr(meta, "report_format", "508").lower()

    doc = Document()

    # -- Title page ----------------------------------------------------------
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("Accessibility Conformance Report")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

    subtitle_map = {
        "508": "Section 508 Edition (VPAT 2.5)",
        "int": "International Edition (EN 301 549)",
        "wcag": "WCAG Edition",
    }
    subtitle = subtitle_map.get(report_format, "WCAG Edition")
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run(subtitle)
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = RGBColor(0x54, 0x54, 0x54)

    doc.add_paragraph()  # spacer

    # -- Metadata table ------------------------------------------------------
    doc.add_heading("Report Information", level=2)
    meta_dict = meta.to_dict()
    meta_rows = [
        ("Product Name", meta_dict.get("product_name", "")),
        ("Company", meta_dict.get("company_name", "")),
        ("URL", meta_dict.get("source_url", "")),
        ("Report Format", report_format.upper()),
        ("WCAG Version", meta_dict.get("wcag_version", "2.2")),
        ("Coverage Level", meta_dict.get("coverage_level", "AA")),
        ("Date", meta_dict.get("created_at", "")),
        ("Review ID", meta_dict.get("review_id", "")),
    ]
    meta_table = doc.add_table(rows=len(meta_rows), cols=2)
    meta_table.style = "Table Grid"
    for idx, (label, value) in enumerate(meta_rows):
        meta_table.rows[idx].cells[0].text = label
        meta_table.rows[idx].cells[1].text = str(value)
        for paragraph in meta_table.rows[idx].cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    doc.add_paragraph()

    # -- Additional metadata sections ----------------------------------------
    product_description = str(meta_dict.get("product_description", "")).strip()
    if product_description:
        doc.add_heading("Product Description", level=3)
        doc.add_paragraph(product_description)

    contact_name = str(meta_dict.get("contact_name", "")).strip()
    contact_email = str(meta_dict.get("contact_email", "")).strip()
    if contact_name or contact_email:
        doc.add_heading("Contact", level=3)
        if contact_name:
            doc.add_paragraph(f"Name: {contact_name}")
        if contact_email:
            doc.add_paragraph(f"Email: {contact_email}")

    notes = str(meta_dict.get("notes", "")).strip()
    if notes:
        doc.add_heading("Notes", level=3)
        doc.add_paragraph(notes)

    evaluation_methods = str(meta_dict.get("evaluation_methods", "")).strip()
    if evaluation_methods:
        doc.add_heading("Evaluation Methods", level=3)
        doc.add_paragraph(evaluation_methods)

    doc.add_paragraph()

    # -- Summary -------------------------------------------------------------
    doc.add_heading("Summary", level=2)
    total = len(results)
    supports = sum(1 for r in results if _conformance_str(r.conformance_level) == ConformanceLevel.SUPPORTS.value)
    partially = sum(1 for r in results if _conformance_str(r.conformance_level) == ConformanceLevel.PARTIALLY_SUPPORTS.value)
    does_not = sum(1 for r in results if _conformance_str(r.conformance_level) == ConformanceLevel.DOES_NOT_SUPPORT.value)
    na_count = sum(1 for r in results if _conformance_str(r.conformance_level) == ConformanceLevel.NOT_APPLICABLE.value)
    ne_count = sum(1 for r in results if _conformance_str(r.conformance_level) == ConformanceLevel.NOT_EVALUATED.value)
    total_findings = sum(len(r.findings) for r in results)

    summary_table = doc.add_table(rows=7, cols=2)
    summary_table.style = "Table Grid"
    summary_items = [
        ("Total Criteria", str(total)),
        ("Supports", str(supports)),
        ("Partially Supports", str(partially)),
        ("Does Not Support", str(does_not)),
        ("Not Applicable", str(na_count)),
        ("Not Evaluated", str(ne_count)),
        ("Total Findings", str(total_findings)),
    ]
    for idx, (label, value) in enumerate(summary_items):
        summary_table.rows[idx].cells[0].text = label
        summary_table.rows[idx].cells[1].text = value
        for paragraph in summary_table.rows[idx].cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    doc.add_paragraph()

    # -- Criteria tables by level --------------------------------------------
    buckets: dict[str, list[dict]] = {"A": [], "AA": [], "AAA": []}
    for r in results:
        key = r.level.upper().strip()
        if key not in buckets:
            key = "A"
        row = r.to_dict()
        row["conformance_display"] = _conformance_str(r.conformance_level)
        buckets[key].append(row)

    _add_criteria_table(doc, "Table 1: Level A", buckets["A"])
    _add_criteria_table(doc, "Table 2: Level AA", buckets["AA"])
    if buckets["AAA"]:
        _add_criteria_table(doc, "Table 3: Level AAA", buckets["AAA"])

    # -- FPC table (Section 508 only) ----------------------------------------
    if report_format == "508":
        fpc_rows = _build_fpc_rows(results)
        if fpc_rows:
            doc.add_heading("Functional Performance Criteria (Section 508)", level=2)
            fpc_table = doc.add_table(rows=1, cols=4)
            fpc_table.style = "Table Grid"
            fpc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr = fpc_table.rows[0].cells
            hdr[0].text = "Code"
            hdr[1].text = "Criteria"
            hdr[2].text = "Conformance Level"
            hdr[3].text = "Remarks and Explanations"
            _style_header_row(fpc_table)
            for item in fpc_rows:
                row = fpc_table.add_row()
                row.cells[0].text = item["code"]
                row.cells[1].text = item["name"]
                _add_conformance_cell(row, 2, item["conformance"])
                row.cells[3].text = item.get("remark", "")
            doc.add_paragraph()

    # -- Findings ------------------------------------------------------------
    all_findings = []
    for r in results:
        for f in r.findings:
            entry = f.to_dict() if hasattr(f, "to_dict") else dict(f)
            entry["criterion_id"] = r.criterion_id
            entry["criterion_name"] = r.criterion_name
            all_findings.append(entry)

    if all_findings:
        sev_order = {"high": 0, "medium": 1, "low": 2, "info": 3}
        all_findings.sort(key=lambda x: (sev_order.get(x.get("severity", "info"), 4), x.get("criterion_id", "")))

        doc.add_heading("Detailed Findings", level=2)
        for finding in all_findings:
            heading_text = f"{finding.get('criterion_id', '')} - {finding.get('issue', 'Finding')}"
            doc.add_heading(heading_text, level=3)
            severity = finding.get("severity", "info")
            sev_color = _SEVERITY_COLORS.get(severity, RGBColor(0x75, 0x75, 0x75))

            sev_para = doc.add_paragraph()
            sev_run = sev_para.add_run(f"Severity: {severity.upper()}")
            sev_run.bold = True
            sev_run.font.color.rgb = sev_color

            if finding.get("element"):
                doc.add_paragraph(f"Element: {finding['element']}")
            if finding.get("location"):
                doc.add_paragraph(f"Location: {finding['location']}")
            if finding.get("css_selector"):
                doc.add_paragraph(f"Selector: {finding['css_selector']}")
            if finding.get("impact"):
                doc.add_paragraph(f"Impact: {finding['impact']}")
            if finding.get("recommendation"):
                doc.add_paragraph(f"Recommendation: {finding['recommendation']}")
            _embed_screenshot(doc, finding.get("screenshot_path"), output_path)

            doc.add_paragraph()  # spacer

    # -- Legal Disclaimer ----------------------------------------------------
    doc.add_paragraph()
    disclaimer_para = doc.add_paragraph()
    disclaimer_run = disclaimer_para.add_run(
        "Legal Disclaimer: This report is provided for informational purposes "
        "only and does not constitute legal advice. Accessibility conformance "
        "was evaluated based on automated and manual testing methods and may "
        "not reflect all potential accessibility barriers. No warranty, express "
        "or implied, is made regarding the completeness or accuracy of this "
        "report."
    )
    disclaimer_run.italic = True
    disclaimer_run.font.size = Pt(9)
    disclaimer_run.font.color.rgb = RGBColor(0x75, 0x75, 0x75)

    # -- Save ----------------------------------------------------------------
    doc.save(str(output_path))
    logger.info("DOCX report written to %s", output_path)
    return str(output_path)
